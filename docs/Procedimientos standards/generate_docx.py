#!/usr/bin/env python3
"""
Converts PRO.003.URM markdown to a professional Word document (.docx).
Uses python-docx to create a well-formatted SIGO CODELCO procedure document.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# --- Configuration ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "PRO.003.URM - Preparacion Rueda de Moldeo - ESTANDAR.md")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "PRO.003.URM - Preparacion Rueda de Moldeo - ESTANDAR.docx")

FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11)
FONT_SIZE_TABLE = Pt(9)
FONT_SIZE_TABLE_HEADER = Pt(9)

HEADER_BG_COLOR = "D6E4F0"  # Light blue
ALT_ROW_COLOR = "F2F2F2"    # Light gray
CODELCO_BLUE = RGBColor(0x1F, 0x47, 0x88)
HEADING_COLOR = RGBColor(0x1F, 0x47, 0x88)


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, top="single", bottom="single", left="single", right="single",
                     color="BFBFBF", sz="4"):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_table_borders(table, color="BFBFBF", sz="4"):
    """Set borders for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text, font_size=None, bold=None):
    """Add text to a paragraph, handling **bold** markdown syntax."""
    # Split text by **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
        run.font.name = FONT_NAME
        if font_size:
            run.font.size = font_size


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_separator_row(line):
    """Check if a markdown table row is a separator (|---|---|)."""
    cells = parse_table_row(line)
    return all(re.match(r'^:?-+:?$', cell.strip()) for cell in cells if cell.strip())


def is_risk_table(headers):
    """Check if this is a risk analysis table (section 4 style)."""
    header_text = ' '.join(headers).upper()
    return 'PELIGROS' in header_text and 'RIESGOS' in header_text


def parse_markdown(filepath):
    """Parse markdown file into structured elements."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    elements = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            elements.append(('hr', None))
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            elements.append(('heading', (level, text)))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped[1:].strip()
            elements.append(('blockquote', text))
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = parse_table_row(table_lines[0])
                rows = []
                for tl in table_lines[1:]:
                    if not is_separator_row(tl):
                        rows.append(parse_table_row(tl))
                elements.append(('table', (headers, rows)))
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            items = []
            while i < len(lines):
                s = lines[i].strip()
                nm = re.match(r'^(\d+)\.\s+(.*)', s)
                if nm:
                    items.append(('num', nm.group(2)))
                    i += 1
                    # Check for sub-items
                    while i < len(lines) and lines[i].strip().startswith('- '):
                        sub = lines[i].strip()[2:]
                        items.append(('sub_bullet', sub))
                        i += 1
                        # Check for sub-sub-items
                        while i < len(lines) and re.match(r'^\s{4,}- ', lines[i]):
                            subsub = lines[i].strip()[2:]
                            items.append(('sub_sub_bullet', subsub))
                            i += 1
                elif s.startswith('- '):
                    # Sub items that are not indented but follow numbered items
                    items.append(('sub_bullet', s[2:]))
                    i += 1
                else:
                    break
            elements.append(('numbered_list', items))
            continue

        # Bullet list
        if stripped.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
                # Check for sub-items (indented bullets)
                while i < len(lines) and re.match(r'^\s{2,}- ', lines[i]):
                    sub_text = lines[i].strip()[2:]
                    items.append(('sub', sub_text))
                    i += 1
            elements.append(('bullet_list', items))
            continue

        # Regular paragraph
        elements.append(('paragraph', stripped))
        i += 1

    return elements


def create_title_page(doc):
    """Create a professional title page."""
    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Main title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROCEDIMIENTO OPERACIONAL")
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.color.rgb = CODELCO_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIGO CODELCO")
    run.font.name = FONT_NAME
    run.font.size = Pt(24)
    run.font.color.rgb = CODELCO_BLUE
    run.bold = True

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = CODELCO_BLUE

    doc.add_paragraph()

    # Document code and title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRO.003.URM")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PREPARACION RUEDA DE MOLDEO")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata block
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, color="1F4788", sz="6")

    meta_data = [
        ("Codigo", "PRO.003.URM"),
        ("Version", "03 - Revision General"),
        ("Fecha", "18/03/2026"),
        ("Unidad", "Refino y Moldeo - Gerencia Fundicion, Division Codelco Chuquicamata"),
        ("Clasificacion", "Actividad Critica del Proceso Productivo"),
    ]

    for idx, (label, value) in enumerate(meta_data):
        row = meta_table.rows[idx]
        # Label cell
        cell_label = row.cells[0]
        cell_label.width = Cm(4)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell_label, "1F4788")

        # Value cell
        cell_value = row.cells[1]
        cell_value.width = Cm(10)
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Approval block
    approval_table = doc.add_table(rows=4, cols=3)
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(approval_table, color="1F4788", sz="4")

    approval_headers = ["Rol", "Nombre", "Cargo"]
    for j, h in enumerate(approval_headers):
        cell = approval_table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F4788")

    approval_data = [
        ("Elaborado por", "Mauricio Pulgar P.", "Ingeniero de Operaciones"),
        ("Aprobado por", "Roderick Cerda S.", "Superintendente Operaciones Fundicion"),
        ("V B", "Bettsy Acosta M.", "Ejecutivo de Negocio GSSO"),
    ]

    for idx, (rol, nombre, cargo) in enumerate(approval_data):
        row = approval_table.rows[idx + 1]
        for j, val in enumerate((rol, nombre, cargo)):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True

    # Page break after title page
    doc.add_page_break()


def add_table_to_doc(doc, headers, rows, is_risk=False):
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = len(headers)

    # Ensure all rows have the correct number of columns
    clean_rows = []
    for row in rows:
        while len(row) < num_cols:
            row.append("")
        clean_rows.append(row[:num_cols])

    table = doc.add_table(rows=1 + len(clean_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    for j, header_text in enumerate(headers):
        cell = header_row.cells[j]
        set_cell_shading(cell, HEADER_BG_COLOR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clean header text of bold markers
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', header_text)
        run = p.add_run(clean_text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE_HEADER
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    # Data rows
    for i, row_data in enumerate(clean_rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # Alternating row shading
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_COLOR)

            p = cell.paragraphs[0]
            # For risk tables, use smaller font and left alignment
            if is_risk:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Center alignment for non-risk tables with short content
                if len(cell_text) < 30:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_formatted_text(p, cell_text, font_size=FONT_SIZE_TABLE)

    # Set column widths for risk tables
    if is_risk and num_cols == 5:
        # ACTIVIDAD | PELIGROS | RIESGOS | CONSECUENCIAS | MEDIDAS DE CONTROL
        widths = [Cm(4.5), Cm(3.5), Cm(3.5), Cm(3.0), Cm(4.5)]
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = width

    return table


def build_document(elements):
    """Build the Word document from parsed elements."""
    doc = Document()

    # --- Page Setup ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- Default font style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY

    # Configure heading styles
    for level, size, color in [
        (1, FONT_SIZE_H1, HEADING_COLOR),
        (2, FONT_SIZE_H2, HEADING_COLOR),
        (3, FONT_SIZE_H3, HEADING_COLOR),
        (4, FONT_SIZE_H4, HEADING_COLOR),
    ]:
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = FONT_NAME
        h_font.size = size
        h_font.bold = True
        h_font.color.rgb = color

    # --- Title Page ---
    create_title_page(doc)

    # --- Add header/footer ---
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("PRO.003.URM - Preparacion Rueda de Moldeo | SIGO CODELCO")
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("PROCEDIMIENTO OPERACIONAL - SIGO CODELCO | Division Codelco Chuquicamata")
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Process Elements ---
    skip_first_heading = True  # Skip the H1 title (already in title page)
    skip_second_heading = True  # Skip "PROCEDIMIENTO OPERACIONAL" (already in title page)
    first_table_skipped = False
    second_table_skipped = False

    for elem_type, elem_data in elements:
        if elem_type == 'heading':
            level, text = elem_data

            # Skip the first H1 and the "PROCEDIMIENTO OPERACIONAL" H2
            if skip_first_heading and level == 1:
                skip_first_heading = False
                continue
            if skip_second_heading and level == 2 and 'PROCEDIMIENTO OPERACIONAL' in text.upper():
                skip_second_heading = False
                continue

            # Clean bold markers from heading text
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_heading(clean_text, level=min(level, 4))

        elif elem_type == 'table':
            headers, rows = elem_data

            # Skip the first two tables (metadata/approval - already in title page)
            if not first_table_skipped:
                first_table_skipped = True
                continue
            if not second_table_skipped:
                second_table_skipped = True
                continue

            risk = is_risk_table(headers)
            add_table_to_doc(doc, headers, rows, is_risk=risk)
            # Add spacing after table
            doc.add_paragraph()

        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            add_formatted_text(p, elem_data, font_size=FONT_SIZE_BODY)

        elif elem_type == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            # Add a left border effect via shading
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>')
            pPr.append(shading)
            add_formatted_text(p, elem_data, font_size=FONT_SIZE_BODY)

        elif elem_type == 'bullet_list':
            for item in elem_data:
                if isinstance(item, tuple) and item[0] == 'sub':
                    p = doc.add_paragraph(style='List Bullet 2')
                    add_formatted_text(p, item[1], font_size=FONT_SIZE_BODY)
                else:
                    p = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(p, item, font_size=FONT_SIZE_BODY)

        elif elem_type == 'numbered_list':
            for item_type, item_text in elem_data:
                if item_type == 'num':
                    p = doc.add_paragraph(style='List Number')
                    add_formatted_text(p, item_text, font_size=FONT_SIZE_BODY)
                elif item_type == 'sub_bullet':
                    p = doc.add_paragraph(style='List Bullet 2')
                    add_formatted_text(p, item_text, font_size=FONT_SIZE_BODY)
                elif item_type == 'sub_sub_bullet':
                    p = doc.add_paragraph(style='List Bullet 3')
                    add_formatted_text(p, item_text, font_size=FONT_SIZE_BODY)

        elif elem_type == 'hr':
            # Add a thin line separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("_" * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)

    return doc


def main():
    print(f"Reading markdown: {INPUT_FILE}")
    if not os.path.exists(INPUT_FILE):
        print(f"ERROR: Input file not found: {INPUT_FILE}")
        return

    elements = parse_markdown(INPUT_FILE)
    print(f"Parsed {len(elements)} elements from markdown")

    print("Building Word document...")
    doc = build_document(elements)

    print(f"Saving to: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print("Document generated successfully!")

    # Verify file size
    size_kb = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()

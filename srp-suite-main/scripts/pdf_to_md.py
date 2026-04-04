"""
Document to Markdown Converter (PDF + Word)
Usage:
    python pdf_to_md.py <input_folder> [output_folder]

Example:
    python pdf_to_md.py "C:\mis_documentos" "C:\markdown_output"

Supports: .pdf, .docx, .doc
If no output folder is provided, creates 'markdown_output' next to input folder.
"""

import sys
import os
import pymupdf4llm
from docx import Document


def convert_pdf_to_md(input_path: str) -> str:
    """Convert a PDF file to Markdown string."""
    return pymupdf4llm.to_markdown(input_path)


def convert_docx_to_md(input_path: str) -> str:
    """Convert a Word (.docx) file to Markdown string."""
    doc = Document(input_path)
    md_lines = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # Remove namespace

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is None:
                continue

            text = para.text.strip()
            if not text:
                md_lines.append("")
                continue

            style_name = para.style.name if para.style else ""

            # Handle headings
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                md_lines.append(f"{'#' * level} {text}")

            # Handle list items
            elif style_name.startswith("List"):
                # Check for numbered list
                if "Number" in style_name or "number" in style_name:
                    md_lines.append(f"1. {text}")
                else:
                    md_lines.append(f"- {text}")

            # Regular paragraph with inline formatting
            else:
                formatted = _format_paragraph_runs(para)
                md_lines.append(formatted)

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    md_lines.append(_table_to_md(table))
                    break

    return "\n\n".join(md_lines)


def _format_paragraph_runs(para) -> str:
    """Format paragraph runs with bold/italic markdown."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts) if parts else para.text


def _table_to_md(table) -> str:
    """Convert a Word table to Markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Build markdown table
    lines = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row if needed
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

    return "\n".join(lines)


def clean_filename(name: str) -> str:
    """Clean a filename for use as markdown file."""
    clean = name.replace(" ", "_")
    for char in "()[]{}!@#$%^&+=":
        clean = clean.replace(char, "")
    return clean


def convert_file(input_path: str, output_folder: str) -> str:
    """Convert a single file (PDF or DOCX) to Markdown."""
    ext = os.path.splitext(input_path)[1].lower()
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    clean_name = clean_filename(base_name)
    output_file = os.path.join(output_folder, f"{clean_name}.md")

    if ext == ".pdf":
        md_text = convert_pdf_to_md(input_path)
    elif ext in (".docx", ".doc"):
        md_text = convert_docx_to_md(input_path)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(md_text)

    return output_file


def convert_folder(input_folder: str, output_folder: str):
    """Convert all PDFs and Word docs in a folder to Markdown."""
    os.makedirs(output_folder, exist_ok=True)

    supported_ext = (".pdf", ".docx", ".doc")
    files = [
        f
        for f in os.listdir(input_folder)
        if os.path.splitext(f)[1].lower() in supported_ext
    ]

    if not files:
        print(f"\n  No PDF or Word files found in: {input_folder}")
        return

    # Count by type
    pdf_count = sum(1 for f in files if f.lower().endswith(".pdf"))
    doc_count = len(files) - pdf_count

    print(f"\n{'='*60}")
    print(f"  Document to Markdown Converter")
    print(f"{'='*60}")
    print(f"  Input:   {input_folder}")
    print(f"  Output:  {output_folder}")
    print(f"  Files:   {len(files)} ({pdf_count} PDF, {doc_count} Word)")
    print(f"{'='*60}\n")

    success = 0
    errors = []

    for i, filename in enumerate(files, 1):
        input_path = os.path.join(input_folder, filename)
        ext = os.path.splitext(filename)[1].upper()
        print(f"  [{i}/{len(files)}] {ext} {filename}...", end=" ")

        try:
            output_file = convert_file(input_path, output_folder)
            size_kb = os.path.getsize(output_file) / 1024
            print(f"OK ({size_kb:.1f} KB)")
            success += 1
        except Exception as e:
            print(f"ERROR: {e}")
            errors.append((filename, str(e)))

    print(f"\n{'='*60}")
    print(f"  RESULTS")
    print(f"{'='*60}")
    print(f"  Converted: {success}/{len(files)}")

    if errors:
        print(f"  Errors:    {len(errors)}")
        for fname, error in errors:
            print(f"    - {fname}: {error}")

    print(f"\n  Output: {output_folder}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) >= 3:
        input_folder = sys.argv[1]
        output_folder = sys.argv[2]
    elif len(sys.argv) == 2:
        input_folder = sys.argv[1]
        output_folder = os.path.join(
            os.path.dirname(input_folder), "markdown_output"
        )
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        input_folder = os.path.join(script_dir, "pdfs")
        output_folder = os.path.join(script_dir, "markdown_output")

    if not os.path.isdir(input_folder):
        print(f"\n  Error: Folder not found: {input_folder}")
        print(f"  Usage: python pdf_to_md.py <input_folder> [output_folder]\n")
        sys.exit(1)

    convert_folder(input_folder, output_folder)
